import os
import re
import tempfile
import logging
from PIL import Image
import PyPDF2
from docx import Document as DocxDocument
from flask import current_app

logger = logging.getLogger(__name__)

def extract_text_from_document(file_path, file_type):
    """
    Extract text from a document using text extraction libraries
    based on the file type.
    
    Args:
        file_path: Path to the document file
        file_type: The file extension (pdf, docx, jpg, etc.)
        
    Returns:
        Extracted text content as a string
    """
    try:
        # For image files - OCR disabled temporarily
        if file_type.lower() in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']:
            logger.info("Image text extraction (OCR) is currently disabled")
            return "Image text extraction not available in this version"
        
        # For PDF files
        elif file_type.lower() == 'pdf':
            return extract_text_from_pdf(file_path)
        
        # For Word documents
        elif file_type.lower() in ['docx', 'doc']:
            return extract_text_from_docx(file_path)
        
        # For text files
        elif file_type.lower() in ['txt', 'csv', 'md', 'json', 'xml', 'html']:
            return extract_text_from_textfile(file_path)
        
        # Unsupported file types
        else:
            logger.warning(f"Text extraction not supported for file type: {file_type}")
            return ""
    
    except Exception as e:
        logger.error(f"Text extraction failed: {str(e)}")
        return f"Error extracting text: {str(e)}"

def extract_text_from_image(image_path):
    """
    Extract text from an image file.
    OCR functionality is disabled in this version.
    """
    return "OCR text extraction is not available in this version"

def extract_text_from_pdf(pdf_path):
    """
    Extract text from a PDF file using PyPDF2.
    OCR for scanned pages is disabled in this version.
    """
    extracted_text = []
    
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                
                # If page has no text (possibly a scanned image)
                if not page_text or len(page_text) < 50:
                    logger.info(f"Page {page_num+1} appears to be a scanned image. OCR is not available.")
                    page_text = f"[Page {page_num+1} may contain scanned content. Text extraction limited.]"
                
                if page_text:
                    extracted_text.append(page_text)
    
    except Exception as e:
        logger.error(f"PDF text extraction failed: {str(e)}")
        extracted_text.append(f"PDF text extraction failed: {str(e)}")
    
    return "\n\n".join(extracted_text)

def extract_text_from_docx(docx_path):
    """
    Extract text from a Word document using python-docx.
    """
    try:
        doc = DocxDocument(docx_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {str(e)}")
        return f"DOCX text extraction failed: {str(e)}"

def extract_text_from_textfile(file_path):
    """
    Extract text from a plain text file.
    """
    try:
        with open(file_path, 'r', errors='ignore') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Text file extraction failed: {str(e)}")
        return f"Text file extraction failed: {str(e)}"
