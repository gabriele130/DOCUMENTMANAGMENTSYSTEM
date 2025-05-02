import os
import re
import datetime
import tempfile
from flask import current_app
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document as DocxDocument
from openpyxl import load_workbook
import magic
import html
import logging
from services.central_storage import save_file_to_central_storage, get_file_from_storage

def allowed_file(filename):
    """Check if the file has an allowed extension."""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in current_app.config['ALLOWED_EXTENSIONS']

def get_file_type(file_path):
    """Detect file type using python-magic."""
    mime = magic.Magic(mime=True)
    file_type = mime.from_file(file_path)
    return file_type

def save_document(file, owner_id):
    """
    Salva un documento utilizzando il sistema di storage centralizzato
    e restituisce i metadati del file salvato.
    
    Args:
        file: L'oggetto file da salvare
        owner_id: ID dell'utente proprietario
        
    Returns:
        dict: Dizionario con i metadati del file salvato
    """
    # Utilizza il sistema di storage centralizzato
    try:
        # Prepara il nome originale del file
        original_filename = secure_filename(file.filename)
        
        # Salva il file nel repository centralizzato
        storage_result = save_file_to_central_storage(
            file_obj=file,
            original_filename=original_filename,
            create_backup=True  # Crea automaticamente un backup
        )
        
        if not storage_result:
            raise Exception("Errore durante il salvataggio del file nel repository centralizzato")
        
        # Estrai il tipo di file dall'estensione
        file_type = original_filename.rsplit('.', 1)[1].lower() if '.' in original_filename else 'unknown'
        
        # Registra l'operazione
        logging.info(f"Documento salvato nel repository centralizzato: {storage_result['file_path']}")
        if storage_result['backup_path']:
            logging.info(f"Backup creato: {storage_result['backup_path']}")
        
        # Restituisci i metadati del file
        return {
            'filename': storage_result['filename'],
            'original_filename': original_filename,
            'file_path': storage_result['file_path'],
            'file_type': file_type,
            'file_size': storage_result['file_size']
        }
    except Exception as e:
        logging.error(f"Errore durante il salvataggio del documento: {str(e)}")
        
        # Fallback al metodo originale se il sistema centralizzato fallisce
        try:
            filename = secure_filename(file.filename)
            unique_filename = f"{owner_id}_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
            
            # Salva il file nella directory principale di upload
            file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], unique_filename)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            file.save(file_path)
            
            # Salva anche una copia nella cache di backup per garantire la persistenza
            cache_path = os.path.join(current_app.config['DOCUMENT_CACHE'], unique_filename)
            os.makedirs(os.path.dirname(cache_path), exist_ok=True)
            import shutil
            shutil.copy2(file_path, cache_path)
            
            file_size = os.path.getsize(file_path)
            file_type = filename.rsplit('.', 1)[1].lower() if '.' in filename else 'unknown'
            
            # Registra entrambi i percorsi per maggiore ridondanza
            logging.info(f"Documento salvato con metodo di fallback in: {file_path} e {cache_path}")
            
            return {
                'filename': unique_filename,
                'original_filename': filename,
                'file_path': file_path,
                'file_type': file_type,
                'file_size': file_size
            }
        except Exception as fallback_error:
            logging.error(f"Anche il metodo di fallback è fallito: {str(fallback_error)}")
            raise Exception(f"Impossibile salvare il file: {str(e)} - Fallback error: {str(fallback_error)}")

def extract_document_metadata(file_path, file_type):
    """Extract metadata from a document file based on its type."""
    metadata = {}
    
    try:
        if file_type == 'pdf':
            with open(file_path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                if pdf.metadata:
                    for key, value in pdf.metadata.items():
                        # Clean the key name
                        clean_key = key.replace('/', '_').lower()
                        # Gestione caratteri NULL nei valori dei metadati
                        str_value = str(value)
                        if '\x00' in str_value:
                            str_value = str_value.replace('\x00', ' ').strip()
                        metadata[clean_key] = str_value
                metadata['page_count'] = len(pdf.pages)
                
        elif file_type in ['docx', 'doc']:
            doc = DocxDocument(file_path)
            metadata['page_count'] = len(doc.paragraphs)
            
            # Extract core properties if available
            core_props = getattr(doc, 'core_properties', None)
            if core_props:
                if core_props.author:
                    author = str(core_props.author)
                    if '\x00' in author:
                        author = author.replace('\x00', ' ').strip()
                    metadata['author'] = author
                if core_props.created:
                    created = str(core_props.created)
                    if '\x00' in created:
                        created = created.replace('\x00', ' ').strip()
                    metadata['created'] = created
                if core_props.modified:
                    modified = str(core_props.modified)
                    if '\x00' in modified:
                        modified = modified.replace('\x00', ' ').strip()
                    metadata['modified'] = modified
                if core_props.title:
                    title = str(core_props.title)
                    if '\x00' in title:
                        title = title.replace('\x00', ' ').strip()
                    metadata['title'] = title
                if core_props.subject:
                    subject = str(core_props.subject)
                    if '\x00' in subject:
                        subject = subject.replace('\x00', ' ').strip()
                    metadata['subject'] = subject
            
        elif file_type in ['xlsx', 'xls']:
            wb = load_workbook(file_path, read_only=True)
            metadata['sheet_count'] = len(wb.sheetnames)
            sheet_names = ', '.join(wb.sheetnames)
            if '\x00' in sheet_names:
                sheet_names = sheet_names.replace('\x00', ' ').strip()
            metadata['sheet_names'] = sheet_names
            
            # Extract document properties if available
            props = wb.properties
            if props:
                if props.creator:
                    creator = str(props.creator)
                    if '\x00' in creator:
                        creator = creator.replace('\x00', ' ').strip()
                    metadata['author'] = creator
                if props.created:
                    created = str(props.created)
                    if '\x00' in created:
                        created = created.replace('\x00', ' ').strip()
                    metadata['created'] = created
                if props.modified:
                    modified = str(props.modified)
                    if '\x00' in modified:
                        modified = modified.replace('\x00', ' ').strip()
                    metadata['modified'] = modified
                if props.title:
                    title = str(props.title)
                    if '\x00' in title:
                        title = title.replace('\x00', ' ').strip()
                    metadata['title'] = title
                if props.subject:
                    subject = str(props.subject)
                    if '\x00' in subject:
                        subject = subject.replace('\x00', ' ').strip()
                    metadata['subject'] = subject
        
        # Add general file metadata
        file_stats = os.stat(file_path)
        metadata['file_size_bytes'] = file_stats.st_size
        metadata['file_size_kb'] = round(file_stats.st_size / 1024, 2)
        metadata['file_size_mb'] = round(file_stats.st_size / (1024 * 1024), 2)
        metadata['created_date'] = datetime.datetime.fromtimestamp(file_stats.st_ctime).strftime('%Y-%m-%d %H:%M:%S')
        metadata['modified_date'] = datetime.datetime.fromtimestamp(file_stats.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
        
    except Exception as e:
        current_app.logger.error(f"Error extracting metadata: {str(e)}")
        metadata['error'] = f"Metadata extraction failed: {str(e)}"
    
    return metadata

def get_document_preview(document):
    """Generate HTML preview content for a document."""
    file_path = document.file_path
    file_type = document.file_type
    preview_html = ""
    
    # Verifica che il file esista nel percorso originale
    if not os.path.exists(file_path):
        logging.warning(f"File non trovato al percorso originale: {file_path}")
        
        # Tenta di recuperare il file tramite il sistema di storage centralizzato
        recovered_path = get_file_from_storage(document.filename, document_id=document.id)
        
        if recovered_path:
            # Se trovato, aggiorna il percorso nel database
            from app import db
            document.file_path = recovered_path
            db.session.commit()
            logging.info(f"Percorso file aggiornato per documento ID: {document.id}")
            file_path = recovered_path
        else:
            # Se non trovato nel sistema centralizzato, prova il metodo precedente
            
            # Estrai nome file dalla parte finale del path (dopo l'ultimo slash o backslash)
            original_file = os.path.basename(file_path)
            base_filename = os.path.basename(document.filename)
            
            # Estrai nome file senza UUID e timestamp (se presente)
            # Ad esempio, da "fdadaadb-2cf1-422d-b5a4-d79640a1dfb6_Atto.pdf" estraiamo "Atto.pdf"
            original_name_parts = original_file.split('_', 1)
            simplified_original = original_name_parts[1] if len(original_name_parts) > 1 else original_file
            
            # Il nome del file potrebbe avere diversi formati a seconda di come è stato caricato
            # Prova tutte le possibili varianti del nome file
            alternative_filenames = [
                document.filename,                  # Nome completo archiviato
                base_filename,                      # Solo nome senza percorso
                document.original_filename,         # Nome originale all'upload
                simplified_original,                # Nome senza UUID
                document.original_filename.replace(' ', '_')  # Con underscore invece di spazi
            ]
            
            # Directory di storage centralizzato
            from services.central_storage import ORIGINAL_FILES_DIR, BACKUP_FILES_DIR
            
            # Elenco di cartelle da esplorare (iniziando con lo storage centralizzato)
            possible_dirs = [
                ORIGINAL_FILES_DIR,                 # Storage centralizzato principale
                BACKUP_FILES_DIR,                   # Storage centralizzato backup
                'document_storage/originals',       # Percorso relativo centralizzato principale
                'document_storage/backup',          # Percorso relativo centralizzato backup
                current_app.config['UPLOAD_FOLDER'],  # Directory principale di upload
                current_app.config['DOCUMENT_CACHE'], # Cache di backup
                os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'uploads'),
                os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'document_cache'),
                'uploads',
                'document_cache',
                'attached_assets',
                os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'attached_assets'),
                '/home/runner/workspace/uploads',
                '/home/runner/workspace/document_cache',
                '/home/runner/workspace/attached_assets',
                os.getcwd(),
                # Aggiungi anche il nuovo sistema centralizzato con percorso assoluto
                '/home/runner/workspace/document_storage/originals',
                '/home/runner/workspace/document_storage/backup'
            ]
            
            # Genera tutte le possibili combinazioni
            alternatives = []
            for dir_path in possible_dirs:
                for filename in alternative_filenames:
                    if filename:
                        alternatives.append(os.path.join(dir_path, filename))
                        
            # Log per debug
            logging.info(f"Cercando file: {document.filename} in {len(alternatives)} percorsi alternativi")
            
            file_found = False
            for alternative_path in alternatives:
                if os.path.exists(alternative_path):
                    # Se trovato, aggiorna il percorso nel db
                    from app import db
                    document.file_path = alternative_path
                    db.session.commit()
                    logging.info(f"Percorso file aggiornato per documento ID: {document.id}, nuovo percorso: {alternative_path}")
                    file_path = alternative_path
                    file_found = True
                    break
            
            if not file_found:
                logging.error(f"File non trovato: {file_path} o alternative: {alternatives}")
                return f"""
                <div class="alert alert-danger">
                    <h4>File non trovato</h4>
                    <p>Il file non è stato trovato nel sistema. Contattare l'amministratore.</p>
                    <p>ID documento: {document.id}, Filename: {document.filename}</p>
                    <p>Percorso registrato: {document.file_path}</p>
                    <p>Nota: È stato creato un sistema di storage centralizzato. Se il problema persiste, un amministratore deve eseguire la migrazione dei documenti al nuovo sistema.</p>
                </div>
                """
    
    try:
        if file_type == 'pdf':
            preview_html = f"""
            <div class="pdf-viewer">
                <iframe src="data:application/pdf;base64,{get_base64_content(file_path)}" 
                    width="100%" height="600px" style="border: none;"></iframe>
            </div>
            """
        
        elif file_type in ['docx', 'doc']:
            # For Word documents, extract text and format as HTML
            doc = DocxDocument(file_path)
            content = []
            for para in doc.paragraphs:
                if para.text:
                    content.append(f"<p>{html.escape(para.text)}</p>")
            preview_html = f"<div class='document-preview'>{' '.join(content)}</div>"
        
        elif file_type in ['xlsx', 'xls']:
            # For Excel, display first sheet as table
            wb = load_workbook(file_path, read_only=True)
            sheet = wb.active
            content = ["<table class='table table-bordered table-striped'>"]
            
            # Limit to first 100 rows for performance
            max_rows = min(100, sheet.max_row)
            max_cols = min(10, sheet.max_column)
            
            for i, row in enumerate(sheet.iter_rows(max_row=max_rows, max_col=max_cols)):
                if i == 0:
                    content.append("<thead><tr>")
                    for cell in row:
                        content.append(f"<th>{html.escape(str(cell.value or ''))}</th>")
                    content.append("</tr></thead><tbody>")
                else:
                    content.append("<tr>")
                    for cell in row:
                        content.append(f"<td>{html.escape(str(cell.value or ''))}</td>")
                    content.append("</tr>")
            
            content.append("</tbody></table>")
            preview_html = ''.join(content)
        
        elif file_type in ['txt']:
            # For text files, just display the content
            with open(file_path, 'r', errors='ignore') as f:
                text_content = f.read(10000)  # Limit to first 10K chars
                preview_html = f"<pre>{html.escape(text_content)}</pre>"
        
        elif file_type in ['jpg', 'jpeg', 'png', 'gif']:
            # For images, use img tag with base64 encoding
            preview_html = f"""
            <div class="image-preview text-center">
                <img src="data:image/{file_type};base64,{get_base64_content(file_path)}" 
                    class="img-fluid" style="max-height: 600px;" alt="{document.original_filename}">
            </div>
            """
        
        else:
            preview_html = f"""
            <div class="alert alert-info">
                <h4>Anteprima non disponibile</h4>
                <p>Il tipo di file '{file_type}' non può essere visualizzato in anteprima. Scarica il file per visualizzarne il contenuto.</p>
            </div>
            """
    
    except Exception as e:
        current_app.logger.error(f"Error generating preview: {str(e)}")
        preview_html = f"""
        <div class="alert alert-danger">
            <h4>Errore anteprima</h4>
            <p>Si è verificato un errore durante la generazione dell'anteprima: {str(e)}</p>
        </div>
        """
    
    return preview_html

def get_base64_content(file_path):
    """Convert file content to base64 string."""
    import base64
    with open(file_path, 'rb') as f:
        return base64.b64encode(f.read()).decode('utf-8')

def sanitize_filename(filename):
    """Remove dangerous characters from a filename."""
    return re.sub(r'[^\w\.-]', '_', filename)
